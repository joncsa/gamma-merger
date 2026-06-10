import os
import re
import base64
import requests
import tempfile
import shutil
import zipfile
from flask import Flask, request, jsonify, send_file
from lxml import etree

app = Flask(__name__)

# XML namespaces
NS_RELS      = 'http://schemas.openxmlformats.org/package/2006/relationships'
NS_CT        = 'http://schemas.openxmlformats.org/package/2006/content-types'
NS_PRES      = 'http://schemas.openxmlformats.org/presentationml/2006/main'
NS_R         = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
NS_P         = 'http://schemas.openxmlformats.org/presentationml/2006/main'

SLIDE_CONTENT_TYPE  = 'application/vnd.openxmlformats-officedocument.presentationml.slide+xml'
SLIDE_REL_TYPE      = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide'

# ── Helpers ───────────────────────────────────────────────────────────────────

def sanitize_filename(name):
    name = re.sub(r'[<>:"/\\|?*&]', '-', name)
    name = re.sub(r'-+', '-', name)
    return name.strip('-')


def get_xml(zf, path):
    """Read and parse an XML file from a ZIP."""
    with zf.open(path) as f:
        return etree.parse(f).getroot()


def set_xml(zf_out, path, root):
    """Write an XML element tree to a ZIP entry."""
    zf_out.writestr(path, etree.tostring(root, xml_declaration=True,
                                          encoding='UTF-8', standalone=True))


# ── OneDrive helpers ──────────────────────────────────────────────────────────

def get_onedrive_access_token():
    client_id     = os.environ['ONEDRIVE_CLIENT_ID']
    client_secret = os.environ['ONEDRIVE_CLIENT_SECRET']
    refresh_token = os.environ['ONEDRIVE_REFRESH_TOKEN']

    response = requests.post(
        'https://login.microsoftonline.com/common/oauth2/v2.0/token',
        data={
            'client_id':     client_id,
            'client_secret': client_secret,
            'refresh_token': refresh_token,
            'grant_type':    'refresh_token',
            'scope':         'offline_access Files.ReadWrite.All',
        },
        timeout=30
    )
    response.raise_for_status()
    token_data = response.json()
    if 'refresh_token' in token_data:
        os.environ['ONEDRIVE_REFRESH_TOKEN'] = token_data['refresh_token']
    return token_data['access_token']


def upload_to_onedrive(file_path, filename, folder_id, access_token):
    CHUNK_SIZE = 10 * 1024 * 1024

    create_session_url = (
        f'https://graph.microsoft.com/v1.0/me/drive/items/{folder_id}:/{filename}:/createUploadSession'
    )
    session_response = requests.post(
        create_session_url,
        headers={
            'Authorization': f'Bearer {access_token}',
            'Content-Type':  'application/json',
        },
        json={'item': {'@microsoft.graph.conflictBehavior': 'rename', 'name': filename}},
        timeout=30
    )
    session_response.raise_for_status()
    upload_url = session_response.json()['uploadUrl']

    file_size = os.path.getsize(file_path)
    uploaded  = 0

    with open(file_path, 'rb') as f:
        while uploaded < file_size:
            chunk         = f.read(CHUNK_SIZE)
            chunk_len     = len(chunk)
            range_end     = uploaded + chunk_len - 1
            content_range = f'bytes {uploaded}-{range_end}/{file_size}'

            chunk_response = requests.put(
                upload_url,
                headers={
                    'Content-Length': str(chunk_len),
                    'Content-Range':  content_range,
                },
                data=chunk,
                timeout=120
            )
            if chunk_response.status_code not in (200, 201, 202):
                chunk_response.raise_for_status()
            uploaded += chunk_len

    result = chunk_response.json()
    return result.get('webUrl', result.get('id', 'uploaded'))


# ── PPTX helpers ──────────────────────────────────────────────────────────────

def download_file(url, dest_path):
    response = requests.get(url, allow_redirects=True, timeout=60)
    response.raise_for_status()
    with open(dest_path, 'wb') as f:
        f.write(response.content)


def merge_pptx_files(pptx_paths, output_path):
    """
    Merge multiple PPTX files at the ZIP/package level.
    Copies slide XML, rels, and media files directly without touching
    their contents — preserving all images including SVG exactly as-is.
    """
    if len(pptx_paths) == 1:
        shutil.copy(pptx_paths[0], output_path)
        return

    # Work on a copy of Part 1 as our base
    base_path = output_path + '.base.pptx'
    shutil.copy(pptx_paths[0], base_path)

    try:
        for src_path in pptx_paths[1:]:
            _append_pptx(base_path, src_path)
        shutil.move(base_path, output_path)
    except Exception:
        if os.path.exists(base_path):
            os.remove(base_path)
        raise


def _append_pptx(base_path, src_path):
    """
    Append all slides from src_path into base_path at the package level.
    Handles filename collisions for slides and media by renaming src entries.
    """
    tmp_path = base_path + '.tmp'

    with zipfile.ZipFile(base_path, 'r') as base_zip, \
         zipfile.ZipFile(src_path,  'r') as src_zip,  \
         zipfile.ZipFile(tmp_path,  'w', zipfile.ZIP_DEFLATED) as out_zip:

        base_names = set(base_zip.namelist())
        src_names  = set(src_zip.namelist())

        # ── 1. Determine existing slide count in base ──────────────────────
        base_slides = sorted([
            n for n in base_names
            if re.match(r'ppt/slides/slide\d+\.xml$', n)
        ], key=lambda x: int(re.search(r'\d+', x.split('/')[-1]).group()))
        base_slide_count = len(base_slides)

        # ── 2. Enumerate slides in src ─────────────────────────────────────
        src_slides = sorted([
            n for n in src_names
            if re.match(r'ppt/slides/slide\d+\.xml$', n)
        ], key=lambda x: int(re.search(r'\d+', x.split('/')[-1]).group()))

        # ── 3. Build media rename map to avoid collisions ──────────────────
        # e.g. ppt/media/image-1-1.png → ppt/media/image-1-1_src.png if collision
        media_rename = {}
        for name in src_names:
            if name.startswith('ppt/media/'):
                if name in base_names:
                    # Rename: insert '_s{N}' before extension
                    base_n, ext = os.path.splitext(name)
                    new_name = base_n + '_s' + str(base_slide_count) + ext
                    media_rename[name] = new_name
                else:
                    media_rename[name] = name

        # ── 4. Build slide rename map ──────────────────────────────────────
        slide_rename = {}
        for i, slide_name in enumerate(src_slides):
            new_num = base_slide_count + i + 1
            new_slide_name = f'ppt/slides/slide{new_num}.xml'
            slide_rename[slide_name] = new_slide_name
            # Also map the .rels file
            old_rels = slide_name.replace('ppt/slides/', 'ppt/slides/_rels/') + '.rels'
            new_rels = new_slide_name.replace('ppt/slides/', 'ppt/slides/_rels/') + '.rels'
            slide_rename[old_rels] = new_rels

        # ── 5. Copy all base files into output ─────────────────────────────
        for name in base_names:
            if name in ('ppt/presentation.xml',
                        'ppt/_rels/presentation.xml.rels',
                        '[Content_Types].xml'):
                continue  # Handle these separately
            out_zip.writestr(name, base_zip.read(name))

        # ── 6. Copy src files into output (with renames) ───────────────────
        skip_src = {
            'ppt/presentation.xml',
            'ppt/_rels/presentation.xml.rels',
            '[Content_Types].xml',
            '_rels/.rels',
            'docProps/app.xml',
            'docProps/core.xml',
        }
        for name in src_names:
            if name in skip_src:
                continue

            data = src_zip.read(name)

            if name in slide_rename:
                # It's a slide XML or slide .rels file — rewrite media refs
                new_name = slide_rename[name]
                if name.endswith('.rels'):
                    data = _rewrite_rels(data, media_rename)
                out_zip.writestr(new_name, data)

            elif name in media_rename:
                out_zip.writestr(media_rename[name], data)

            elif name.startswith('ppt/slides/'):
                # Already handled above via slide_rename
                pass

            elif name.startswith('ppt/media/'):
                # Already handled above via media_rename
                pass

            else:
                # Other src parts (slideLayouts, slideMasters, theme, etc.)
                # Only copy if not already in base to avoid conflicts
                if name not in base_names:
                    out_zip.writestr(name, data)

        # ── 7. Update presentation.xml — add new sldIdLst entries ─────────
        pres_root = etree.fromstring(base_zip.read('ppt/presentation.xml'))
        sldIdLst  = pres_root.find(f'{{{NS_P}}}sldIdLst')
        if sldIdLst is None:
            sldIdLst = etree.SubElement(pres_root, f'{{{NS_P}}}sldIdLst')

        # Find max existing sldId
        existing_ids = [
            int(el.get('id', 0))
            for el in sldIdLst.findall(f'{{{NS_P}}}sldId')
        ]
        max_id = max(existing_ids) if existing_ids else 255

        # Find max existing rId in presentation rels
        pres_rels_root = etree.fromstring(
            base_zip.read('ppt/_rels/presentation.xml.rels')
        )
        existing_rids = [
            int(re.sub(r'\D', '', el.get('Id', 'rId0')))
            for el in pres_rels_root.findall(f'{{{NS_RELS}}}Relationship')
            if re.sub(r'\D', '', el.get('Id', ''))
        ]
        max_rid = max(existing_rids) if existing_rids else 0

        for i, src_slide in enumerate(src_slides):
            new_num   = base_slide_count + i + 1
            new_rid   = f'rId{max_rid + new_num}'
            new_id    = max_id + i + 1
            new_slide = f'ppt/slides/slide{new_num}.xml'

            # Add sldId entry to presentation.xml
            sld_id_el = etree.SubElement(sldIdLst, f'{{{NS_P}}}sldId')
            sld_id_el.set('id',          str(new_id))
            sld_id_el.set(f'{{{NS_R}}}id', new_rid)

            # Add relationship to presentation.xml.rels
            rel_el = etree.SubElement(
                pres_rels_root, f'{{{NS_RELS}}}Relationship'
            )
            rel_el.set('Id',     new_rid)
            rel_el.set('Type',   SLIDE_REL_TYPE)
            rel_el.set('Target', f'slides/slide{new_num}.xml')

        out_zip.writestr(
            'ppt/presentation.xml',
            etree.tostring(pres_root, xml_declaration=True,
                           encoding='UTF-8', standalone=True)
        )
        out_zip.writestr(
            'ppt/_rels/presentation.xml.rels',
            etree.tostring(pres_rels_root, xml_declaration=True,
                           encoding='UTF-8', standalone=True)
        )

        # ── 8. Update [Content_Types].xml ─────────────────────────────────
        ct_root = etree.fromstring(base_zip.read('[Content_Types].xml'))

        # Collect existing Override PartNames
        existing_parts = {
            el.get('PartName')
            for el in ct_root.findall(f'{{{NS_CT}}}Override')
        }

        # Add Override entries for new slides
        for i, src_slide in enumerate(src_slides):
            new_num  = base_slide_count + i + 1
            partname = f'/ppt/slides/slide{new_num}.xml'
            if partname not in existing_parts:
                el = etree.SubElement(ct_root, f'{{{NS_CT}}}Override')
                el.set('PartName',    partname)
                el.set('ContentType', SLIDE_CONTENT_TYPE)

        # Add Default entries for any new media extensions (e.g. svg)
        existing_exts = {
            el.get('Extension')
            for el in ct_root.findall(f'{{{NS_CT}}}Default')
        }
        svg_content_type = 'image/svg+xml'
        if 'svg' not in existing_exts:
            # Check if any src media is svg
            has_svg = any(
                n.lower().endswith('.svg')
                for n in src_names
                if n.startswith('ppt/media/')
            )
            if has_svg:
                el = etree.SubElement(ct_root, f'{{{NS_CT}}}Default')
                el.set('Extension',   'svg')
                el.set('ContentType', svg_content_type)

        out_zip.writestr(
            '[Content_Types].xml',
            etree.tostring(ct_root, xml_declaration=True,
                           encoding='UTF-8', standalone=True)
        )

    # Replace base with tmp
    os.replace(tmp_path, base_path)


def _rewrite_rels(data, media_rename):
    """
    Rewrite Target attributes in a slide .rels file
    to reflect renamed media files.
    """
    root = etree.fromstring(data)
    for rel in root.findall(f'{{{NS_RELS}}}Relationship'):
        target = rel.get('Target', '')
        # Target is relative like ../media/image-1-1.png
        # Resolve to full path for lookup
        full = 'ppt/media/' + target.split('../media/')[-1]
        if full in media_rename and media_rename[full] != full:
            new_basename = os.path.basename(media_rename[full])
            rel.set('Target', '../media/' + new_basename)
    return etree.tostring(root, xml_declaration=True,
                          encoding='UTF-8', standalone=True)


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok', 'service': 'gamma-merger'})


@app.route('/merge', methods=['POST'])
def merge():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No JSON body provided'}), 400

    urls  = data.get('urls',  [])
    files = data.get('files', [])
    if not urls and not files:
        return jsonify({'error': 'Provide urls or files array'}), 400

    work_dir   = tempfile.mkdtemp()
    pptx_paths = []

    try:
        if urls:
            for i, url in enumerate(urls):
                dest = os.path.join(work_dir, f'part{i+1}.pptx')
                try:
                    download_file(url, dest)
                    pptx_paths.append(dest)
                except Exception as e:
                    return jsonify({'error': f'Failed to download part {i+1}', 'details': str(e)}), 500
        elif files:
            for i, b64 in enumerate(files):
                dest = os.path.join(work_dir, f'part{i+1}.pptx')
                try:
                    with open(dest, 'wb') as f:
                        f.write(base64.b64decode(b64))
                    pptx_paths.append(dest)
                except Exception as e:
                    return jsonify({'error': f'Failed to decode part {i+1}', 'details': str(e)}), 500

        if not pptx_paths:
            return jsonify({'error': 'No valid PPTX files to merge'}), 400

        output_path = os.path.join(work_dir, 'merged.pptx')
        merge_pptx_files(pptx_paths, output_path)

        return send_file(
            output_path,
            mimetype='application/vnd.openxmlformats-officedocument.presentationml.presentation',
            as_attachment=True,
            download_name='presentation_merged.pptx'
        )
    except Exception as e:
        return jsonify({'error': 'Merge failed', 'details': str(e)}), 500
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)


@app.route('/merge-and-upload', methods=['POST'])
def merge_and_upload():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No JSON body provided'}), 400

    urls      = data.get('urls', [])
    raw_name  = data.get('filename', 'merged_presentation.pptx')
    folder_id = os.environ.get('ONEDRIVE_FOLDER_ID', '')

    if not urls:
        return jsonify({'error': 'Provide urls array'}), 400
    if not folder_id:
        return jsonify({'error': 'ONEDRIVE_FOLDER_ID environment variable not set'}), 500

    filename = sanitize_filename(raw_name)
    if not filename.lower().endswith('.pptx'):
        filename += '.pptx'

    work_dir   = tempfile.mkdtemp()
    pptx_paths = []

    try:
        for i, url in enumerate(urls):
            dest = os.path.join(work_dir, f'part{i+1}.pptx')
            try:
                download_file(url, dest)
                pptx_paths.append(dest)
            except Exception as e:
                return jsonify({'error': f'Failed to download part {i+1}', 'details': str(e)}), 500

        if not pptx_paths:
            return jsonify({'error': 'No valid PPTX files to merge'}), 400

        output_path = os.path.join(work_dir, filename)
        merge_pptx_files(pptx_paths, output_path)

        try:
            access_token = get_onedrive_access_token()
        except Exception as e:
            return jsonify({'error': 'Failed to get OneDrive access token', 'details': str(e)}), 500

        try:
            web_url = upload_to_onedrive(output_path, filename, folder_id, access_token)
        except Exception as e:
            return jsonify({'error': 'Failed to upload to OneDrive', 'details': str(e)}), 500

        return jsonify({
            'status':   'uploaded',
            'filename': filename,
            'web_url':  web_url,
        })

    except Exception as e:
        return jsonify({'error': 'merge-and-upload failed', 'details': str(e)}), 500
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)

@app.route('/generate-docx', methods=['POST'])
def generate_docx():
    from docx import Document as DocxDocument
    from docx.shared import Pt
    
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No JSON body provided'}), 400

    content   = data.get('content', '')
    raw_name  = data.get('filename', 'output.docx')
    folder_id = os.environ.get('ONEDRIVE_FOLDER_ID', '')

    if not content:
        return jsonify({'error': 'content is required'}), 400
    if not folder_id:
        return jsonify({'error': 'ONEDRIVE_FOLDER_ID not set'}), 500

    filename = sanitize_filename(raw_name)
    if not filename.lower().endswith('.docx'):
        filename += '.docx'

    work_dir = tempfile.mkdtemp()
    try:
        doc = DocxDocument()
        for line in content.split('\n'):
            para = doc.add_paragraph()
            run  = para.add_run(line)
            run.font.size = Pt(11)

        output_path = os.path.join(work_dir, filename)
        doc.save(output_path)

        try:
            access_token = get_onedrive_access_token()
        except Exception as e:
            return jsonify({'error': 'Failed to get OneDrive token', 'details': str(e)}), 500

        try:
            web_url = upload_to_onedrive(output_path, filename, folder_id, access_token)
        except Exception as e:
            return jsonify({'error': 'Failed to upload to OneDrive', 'details': str(e)}), 500

        return jsonify({
            'status':   'uploaded',
            'filename': filename,
            'web_url':  web_url,
        })

    except Exception as e:
        return jsonify({'error': 'generate-docx failed', 'details': str(e)}), 500
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)
        
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port)
